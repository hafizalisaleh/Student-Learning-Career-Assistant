"""
Resume parser for PDF and DOCX files with AI-powered extraction
"""
import PyPDF2
import pdfplumber
from docx import Document
from typing import Dict, Any, List
import re
import json
from utils.gemini_client import gemini_client

class ResumeParser:
    """Parser for extracting structured data from resumes using AI"""
    
    def __init__(self):
        self.gemini = gemini_client
    
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Parse PDF resume with AI enhancement
        Uses PyPDF2 first, falls back to pdfplumber if that fails

        Args:
            file_path: Path to PDF file

        Returns:
            Parsed resume data with AI extraction
        """
        text = ""

        # Try PyPDF2 first
        try:
            text = self._parse_pdf_pypdf2(file_path)
        except Exception as e:
            print(f"PyPDF2 failed: {e}, trying pdfplumber...")
            # Fallback to pdfplumber
            try:
                text = self._parse_pdf_pdfplumber(file_path)
            except Exception as e2:
                raise ValueError(f"Could not read PDF file. The file may be corrupted or in an unsupported format. Please try converting to DOCX or a different PDF.")

        # Check if we extracted any text
        if not text or len(text.strip()) < 50:
            raise ValueError("Could not extract text from PDF. The PDF may be scanned/image-based. Please upload a text-based PDF or DOCX file.")

        return self._extract_structured_data(text)

    def _parse_pdf_pypdf2(self, file_path: str) -> str:
        """Parse PDF using PyPDF2"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                try:
                    pdf_reader.decrypt('')  # Try empty password
                except Exception:
                    raise ValueError("PDF is password protected. Please upload an unprotected PDF.")

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        return text

    def _parse_pdf_pdfplumber(self, file_path: str) -> str:
        """Parse PDF using pdfplumber (fallback parser)"""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

                # Also try to extract text from tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            row_text = ' '.join([cell or '' for cell in row])
                            if row_text.strip():
                                text += row_text + "\n"

        return text
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX resume with AI enhancement

        Args:
            file_path: Path to DOCX file

        Returns:
            Parsed resume data with AI extraction
        """
        text = ""

        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Also extract text from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

        except Exception as e:
            error_msg = str(e).lower()
            if 'not a valid docx' in error_msg or 'package not found' in error_msg:
                raise ValueError("Invalid DOCX file. Please ensure you're uploading a valid Word document (.docx format).")
            raise ValueError(f"Error reading DOCX file: {str(e)}")

        # Check if we extracted any text
        if not text or len(text.strip()) < 50:
            raise ValueError("Could not extract text from document. The file may be empty or corrupted.")

        return self._extract_structured_data(text)
    
    def _extract_structured_data(self, text: str) -> Dict[str, Any]:
        """
        Extract structured data from resume text using AI

        Args:
            text: Raw resume text

        Returns:
            Comprehensive structured resume data
        """
        # First do rule-based extraction for basic info
        basic_info = self._rule_based_extraction(text)

        # Then enhance with AI extraction
        try:
            ai_data = self._ai_powered_extraction(text)
            # Merge AI results with rule-based
            if ai_data:
                basic_info.update(ai_data)
        except Exception as e:
            print(f"AI extraction failed, using rule-based only: {e}")
            # Add fallback skills extraction
            basic_info['skills'] = self._extract_skills_fallback(text)
            basic_info['name'] = self._extract_name_fallback(text)

        # Ensure required fields exist
        if 'skills' not in basic_info or not basic_info['skills']:
            basic_info['skills'] = self._extract_skills_fallback(text)
        if 'education' not in basic_info:
            basic_info['education'] = []
        if 'experience' not in basic_info:
            basic_info['experience'] = []
        if 'projects' not in basic_info:
            basic_info['projects'] = []

        return basic_info

    def _extract_skills_fallback(self, text: str) -> List[str]:
        """Fallback skill extraction using keyword matching"""
        common_skills = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php',
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django',
            'flask', 'spring', 'sql', 'mysql', 'postgresql', 'mongodb', 'redis',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'git', 'linux',
            'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas',
            'numpy', 'data analysis', 'excel', 'tableau', 'power bi',
            'agile', 'scrum', 'jira', 'communication', 'leadership', 'problem solving'
        ]

        text_lower = text.lower()
        found_skills = []

        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill.title() if len(skill) > 3 else skill.upper())

        return found_skills[:20]  # Limit to 20 skills

    def _extract_name_fallback(self, text: str) -> str:
        """Fallback name extraction - first line usually contains name"""
        lines = text.strip().split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            # Name is usually short and contains only letters/spaces
            if line and len(line) < 50 and re.match(r'^[A-Za-z\s\.\-]+$', line):
                # Skip common section headers
                if line.lower() not in ['resume', 'cv', 'curriculum vitae', 'summary', 'objective']:
                    return line
        return "Unknown"
    
    def _rule_based_extraction(self, text: str) -> Dict[str, Any]:
        """Basic rule-based extraction for contact info"""
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        # Extract phone numbers
        phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,5}[-\s\.]?[0-9]{1,5}'
        phones = re.findall(phone_pattern, text)
        
        # Extract experience years
        experience_pattern = r'(\d+)[\+]?\s*(?:years?|yrs?)'
        experience_matches = re.findall(experience_pattern, text.lower())
        total_experience = max([int(exp) for exp in experience_matches]) if experience_matches else 0
        
        return {
            'raw_text': text,
            'email': emails[0] if emails else None,
            'phone': phones[0] if phones else None,
            'experience_years': total_experience,
            'sections': self._identify_sections(text)
        }
    
    def _ai_powered_extraction(self, text: str) -> Dict[str, Any]:
        """
        Use Gemini AI to extract comprehensive resume data
        """
        prompt = f"""
Analyze this resume and extract detailed structured information in JSON format:

RESUME TEXT:
{text[:4000]}

Extract the following information in this exact JSON structure:
{{
    "name": "full name of candidate",
    "summary": "professional summary or objective (if present)",
    "skills": ["skill1", "skill2", "skill3", ...],
    "technical_skills": ["technical skill1", "technical skill2", ...],
    "soft_skills": ["soft skill1", "soft skill2", ...],
    "education": [
        {{
            "degree": "degree name",
            "institution": "university/college name",
            "year": "graduation year",
            "field": "field of study"
        }}
    ],
    "experience": [
        {{
            "title": "job title",
            "company": "company name",
            "duration": "duration (e.g., 2020-2022)",
            "responsibilities": ["responsibility1", "responsibility2", ...],
            "achievements": ["achievement1", "achievement2", ...]
        }}
    ],
    "projects": [
        {{
            "name": "project name",
            "description": "brief description",
            "technologies": ["tech1", "tech2", ...],
            "achievements": "key achievements"
        }}
    ],
    "certifications": [
        {{
            "name": "certification name",
            "issuer": "issuing organization",
            "year": "year obtained"
        }}
    ],
    "languages": ["language1", "language2", ...],
    "achievements": ["notable achievement1", "achievement2", ...],
    "has_linkedin": true/false,
    "has_github": true/false,
    "has_portfolio": true/false
}}

Important:
- Extract ALL skills (technical and soft skills separately)
- List ALL work experiences with detailed responsibilities
- Include ALL projects with technologies used
- List ALL certifications
- If any section is not found, use empty array []
- Return ONLY valid JSON, no additional text
"""
        
        try:
            response = self.gemini.generate_text(prompt, temperature=0.2)
            
            # Extract JSON from response
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            
            if json_start != -1 and json_end > json_start:
                json_str = response[json_start:json_end]
                ai_data = json.loads(json_str)
                
                # Combine all skills
                all_skills = list(set(
                    ai_data.get('skills', []) + 
                    ai_data.get('technical_skills', []) +
                    ai_data.get('soft_skills', [])
                ))
                
                ai_data['skills'] = all_skills
                return ai_data
            
            return {}
            
        except Exception as e:
            print(f"AI extraction parsing error: {e}")
            return {}
    
    def _identify_sections(self, text: str) -> List[str]:
        """
        Identify resume sections
        
        Args:
            text: Resume text
            
        Returns:
            List of identified sections
        """
        section_keywords = [
            'summary', 'objective', 'experience', 'education', 
            'skills', 'projects', 'certifications', 'achievements',
            'publications', 'references'
        ]
        
        identified_sections = []
        text_lower = text.lower()
        
        for section in section_keywords:
            if section in text_lower:
                identified_sections.append(section.title())
        
        return identified_sections

resume_parser = ResumeParser()
