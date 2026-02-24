"""
Notes API endpoints
"""
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.orm import Session
from config.database import get_db
from notes.models import Note
from notes.schemas import NoteCreate, StudyNoteCreate, NoteUpdate, NoteResponse
from documents.models import Document, ProcessingStatus
from users.auth import get_current_user
from users.models import User
from notes.generator import notes_generator
from core.rag_retriever import rag_retriever
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from utils.logger import logger

router = APIRouter(prefix="/api/notes", tags=["notes"])

@router.post("/generate", response_model=NoteResponse, status_code=status.HTTP_201_CREATED)
async def generate_notes(
    note_data: NoteCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Generate AI-powered notes from document using RAG when available.
    Uses vector similarity search to retrieve relevant chunks,
    falls back to full text extraction if embeddings not available.

    Args:
        note_data: Note creation data
        current_user: Current authenticated user
        db: Database session

    Returns:
        Generated notes
    """
    try:
        # Check if document exists and belongs to user
        document = db.query(Document).filter(
            Document.id == note_data.document_id,
            Document.user_id == current_user.id
        ).first()

        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )

        logger.info(f"Generating notes for document {document.id} by user {current_user.email}")

        # Use RAG retriever to get content (uses embeddings if available, else full text)
        retrieval_result = rag_retriever.get_content_for_generation(
            document=document,
            task_type="notes",
            chunk_count=5
        )

        content = retrieval_result.get("content")
        content_source = retrieval_result.get("source")

        logger.info(f"Content retrieved via {content_source}, chunks_used={retrieval_result.get('chunks_used', 0)}")

        if not content or retrieval_result.get("source") == "error":
            error_msg = retrieval_result.get("error", "Could not extract content from document")
            logger.error(f"Content retrieval failed: {error_msg}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error_msg
            )

        logger.info(f"Content extracted successfully, length: {len(content)} characters")
        
        # Generate notes using Gemini AI
        try:
            notes_content = notes_generator.generate_notes(
                content=content,
                title=note_data.title,
                note_type=note_data.note_type,
                additional_context=note_data.additional_context
            )
            logger.info(f"Notes generated successfully, length: {len(notes_content)} characters")
        except Exception as gen_error:
            logger.error(f"Note generation error: {gen_error}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate notes: {str(gen_error)}"
            )
        
        # Save notes to database
        new_note = Note(
            user_id=current_user.id,
            document_id=note_data.document_id,
            title=note_data.title,
            note_type=note_data.note_type,
            content=notes_content,
            tags=note_data.tags
        )
        
        db.add(new_note)
        db.commit()
        db.refresh(new_note)
        
        logger.info(f"Notes saved successfully with ID: {new_note.id}")
        
        return NoteResponse.from_orm(new_note)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in generate_notes: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An unexpected error occurred: {str(e)}"
        )

@router.post("/study", response_model=NoteResponse, status_code=status.HTTP_201_CREATED)
async def create_study_note(
    note_data: StudyNoteCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a manual study note with BlockNote JSON content"""
    document = db.query(Document).filter(
        Document.id == note_data.document_id,
        Document.user_id == current_user.id
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    new_note = Note(
        user_id=current_user.id,
        document_id=note_data.document_id,
        title=note_data.title,
        note_type='study',
        content=note_data.content,
        content_format='blocknote',
        tags=note_data.tags
    )

    db.add(new_note)
    db.commit()
    db.refresh(new_note)

    logger.info(f"Study note created: {new_note.id} for document {document.id}")

    return NoteResponse.from_orm(new_note)

@router.get("/", response_model=list[NoteResponse])
def get_all_notes(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get all notes for the current user
    
    Args:
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        List of all user's notes
    """
    notes = db.query(Note).filter(
        Note.user_id == current_user.id
    ).order_by(Note.generated_at.desc()).all()
    
    return [NoteResponse.from_orm(note) for note in notes]

@router.get("/document/{document_id}", response_model=list[NoteResponse])
def get_notes_by_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get all notes for a document
    
    Args:
        document_id: Document ID
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        List of notes
    """
    notes = db.query(Note).filter(
        Note.document_id == document_id,
        Note.user_id == current_user.id
    ).all()
    
    return [NoteResponse.from_orm(note) for note in notes]

@router.get("/{note_id}", response_model=NoteResponse)
def get_note(
    note_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a single note by ID"""
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()

    if not note:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Note not found"
        )

    return NoteResponse.from_orm(note)

@router.put("/{note_id}", response_model=NoteResponse)
def update_note(
    note_id: str,
    note_data: NoteUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Update a note's title, content, or tags.

    Args:
        note_id: Note ID (UUID)
        note_data: Fields to update
        current_user: Current authenticated user
        db: Database session

    Returns:
        Updated note
    """
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()

    if not note:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Note not found"
        )

    if note_data.title is not None:
        note.title = note_data.title
    if note_data.content is not None:
        note.content = note_data.content
    if note_data.tags is not None:
        note.tags = note_data.tags
    if note_data.content_format is not None:
        note.content_format = note_data.content_format

    db.commit()
    db.refresh(note)

    logger.info(f"Note {note_id} updated by user {current_user.email}")

    return NoteResponse.from_orm(note)

@router.delete("/{note_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_note(
    note_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Delete a note
    
    Args:
        note_id: Note ID (UUID)
        current_user: Current authenticated user
        db: Database session
    """
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()
    
    if not note:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Note not found"
        )
    
    db.delete(note)
    db.commit()
    return {"message": "Note deleted successfully"}

@router.get("/{note_id}/export/markdown")
async def export_note_to_markdown(
    note_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export note to Markdown format (.md file) for download
    Preserves all formatting including LaTeX formulas, tables, etc.

    Args:
        note_id: Note ID (UUID)
        current_user: Current authenticated user
        db: Database session

    Returns:
        Markdown file response
    """
    try:
        logger.info(f"Markdown export requested for note {note_id} by user {current_user.email}")

        # Get the note
        note = db.query(Note).filter(
            Note.id == note_id,
            Note.user_id == current_user.id
        ).first()

        if not note:
            logger.warning(f"Note {note_id} not found for user {current_user.email}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Note not found"
            )

        # Build markdown content with metadata header
        created_date = note.generated_at.strftime('%Y-%m-%d %H:%M') if note.generated_at else 'N/A'

        markdown_content = f"""# {note.title}

> **Note Type:** {note.note_type.capitalize()} | **Generated:** {created_date}

---

{note.content}
"""

        # Generate filename
        filename = f"{note.title.replace(' ', '_')}.md"

        # Create buffer
        buffer = io.BytesIO(markdown_content.encode('utf-8'))
        buffer.seek(0)

        logger.info(f"Markdown generated successfully for note {note_id}, size: {len(markdown_content)} bytes")

        # Return Markdown response
        return StreamingResponse(
            buffer,
            media_type="text/markdown; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Markdown export error for note {note_id}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to export Markdown: {str(e)}"
        )


@router.get("/{note_id}/export/docx")
async def export_note_to_docx(
    note_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export note to DOCX format (Word document) for download
    Preserves markdown formatting beautifully in Word
    
    Args:
        note_id: Note ID (UUID)
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        DOCX file response
    """
    try:
        logger.info(f"DOCX export requested for note {note_id} by user {current_user.email}")
        
        # Get the note
        note = db.query(Note).filter(
            Note.id == note_id,
            Note.user_id == current_user.id
        ).first()
        
        if not note:
            logger.warning(f"Note {note_id} not found for user {current_user.email}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Note not found"
            )
        
        # Create DOCX document in memory
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(note.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(30, 64, 175)  # Blue color
        
        # Add metadata
        created_date = note.generated_at.strftime('%Y-%m-%d %H:%M') if note.generated_at else 'N/A'
        meta_paragraph = doc.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_paragraph.add_run(f"Type: {note.note_type.capitalize()} | Generated: {created_date}")
        meta_run.font.size = Pt(10)
        meta_run.italic = True
        
        doc.add_paragraph()  # Spacer
        
        # Parse markdown and add to document
        lines = note.content.split('\n')
        current_list = None
        list_level = 0
        
        for line in lines:
            line = line.rstrip()
            
            if not line:
                doc.add_paragraph()  # Empty line
                current_list = None
                continue
            
            # Headings
            if line.startswith('# '):
                current_list = None
                heading = doc.add_heading(line[2:], level=1)
                heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
            elif line.startswith('## '):
                current_list = None
                heading = doc.add_heading(line[3:], level=2)
                heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
            elif line.startswith('### '):
                current_list = None
                heading = doc.add_heading(line[4:], level=3)
                heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
            
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                current_list = None
                text = line[2:]
                # Remove markdown bold/italic
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                paragraph = doc.add_paragraph(text, style='List Bullet')
                
            # Numbered lists
            elif re.match(r'^\d+\.\s', line):
                current_list = None
                text = re.sub(r'^\d+\.\s', '', line)
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                paragraph = doc.add_paragraph(text, style='List Number')
            
            # Tables (simple detection)
            elif line.startswith('|'):
                current_list = None
                # Skip table separators
                if not re.match(r'^\|[\s\-:]+\|', line):
                    # This is table data - for now add as paragraph
                    # Full table parsing would require more complex logic
                    doc.add_paragraph(line)
            
            # Regular paragraphs
            else:
                current_list = None
                paragraph = doc.add_paragraph()
                
                # Parse inline markdown (bold, italic, code)
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
                
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = paragraph.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = paragraph.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.color.rgb = RGBColor(220, 38, 38)
                    else:
                        paragraph.add_run(part)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename
        filename = f"{note.title.replace(' ', '_')}.docx"
        
        logger.info(f"DOCX generated successfully for note {note_id}, size: {buffer.getbuffer().nbytes} bytes")
        
        # Return DOCX response
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX export error for note {note_id}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to export DOCX: {str(e)}"
        )
